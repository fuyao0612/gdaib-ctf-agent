from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from win32com.client import DispatchEx

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "submission" / "deliverables" / "documents"


BLUE = RGBColor(47, 93, 153)
def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def add_table(document: Document, rows: list[list[str]]) -> None:
    width = 6.5 / max(1, len(rows[0]))
    table = document.add_table(rows=0, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].width = Inches(width)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(value.replace("`", ""))
            set_run_font(run, 9, row_index == 0)
            if row_index == 0:
                set_cell_shading(cells[index], "E8EEF5")
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_markdown(document: Document, source: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.rstrip()
        if line.startswith("```"):
            in_code = not in_code
            index += 1
            continue
        if not line:
            document.add_paragraph()
            index += 1
            continue

        if line.startswith("|") and index + 1 < len(lines) and re.match(
            r"^\s*\|(?:\s*:?-+:?\s*\|)+\s*$", lines[index + 1]
        ):
            rows = [[item.strip() for item in line.strip("|").split("|")]]
            index += 2
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                rows.append([item.strip() for item in lines[index].strip().strip("|").split("|")])
                index += 1
            add_table(document, rows)
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            paragraph = document.add_heading(heading.group(2), level=len(heading.group(1)))
        elif re.match(r"^[-*]\s+", line):
            paragraph = document.add_paragraph(re.sub(r"^[-*]\s+", "", line), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            paragraph = document.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
        else:
            paragraph = document.add_paragraph(line.replace("`", ""))

        for run in paragraph.runs:
            set_run_font(run, 9.5 if in_code else None)
            if in_code:
                run.font.name = "Consolas"
        index += 1


def build_report(title: str, subtitle: str, sources: list[str], stem: str) -> Path:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, before, after in (
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ):
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = BLUE
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_before = Pt(145)
    title_run = title_paragraph.add_run(title)
    set_run_font(title_run, 28, True)

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_paragraph.add_run(f"{subtitle}\n版本日期：{date.today().isoformat()}")
    set_run_font(subtitle_run, 14)

    document.add_page_break()
    for source in sources:
        add_markdown(document, ROOT / source)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("御网智元 | 初审提交材料 | 2026-09-03")
    set_run_font(footer_run, 8)

    path = OUTPUT / f"{stem}.docx"
    document.save(path)
    return path


def export_pdfs(paths: list[Path]) -> None:
    word = DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        for path in paths:
            document = word.Documents.Open(str(path.resolve()), ReadOnly=True)
            try:
                document.ExportAsFixedFormat(str(path.with_suffix(".pdf").resolve()), 17)
            finally:
                document.Close(False)
    finally:
        word.Quit()


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    reports = [
        build_report(
            "御网智元技术报告",
            "具备自主决策能力的通用网络安全智能体",
            [
                "docs/submission/source/01-项目摘要.md",
                "docs/submission/source/02-技术方案.md",
                "docs/submission/source/03-系统架构与模块说明.md",
                "docs/submission/source/04-自主决策与可解释性设计.md",
                "docs/submission/source/05-工具协同与扩展机制.md",
                "docs/submission/source/06-安全设计与威胁模型.md",
                "docs/submission/source/11-创新点与同类方案对比.md",
                "docs/submission/source/16-方案设计文档.md",
                "docs/submission/source/10-开发文档.md",
            ],
            "御网智元-技术报告",
        ),
        build_report(
            "御网智元测试与评测报告",
            "质量门禁、通用评测与可复现证据",
            ["docs/submission/source/07-测试与评测报告.md"],
            "御网智元-测试与评测报告",
        ),
        build_report(
            "御网智元方案设计文档",
            "自主决策安全智能体的场景、流程与验证设计",
            ["docs/submission/source/16-方案设计文档.md"],
            "御网智元-方案设计文档",
        ),
        build_report(
            "御网智元开发文档",
            "代码组织、扩展流程与本地开发规范",
            ["docs/submission/source/10-开发文档.md"],
            "御网智元-开发文档",
        ),
        build_report(
            "御网智元用户手册",
            "部署后的任务创建、运行控制与结果审计",
            ["docs/submission/source/09-用户手册.md"],
            "御网智元-用户手册",
        ),
    ]
    export_pdfs(reports)


if __name__ == "__main__":
    main()
